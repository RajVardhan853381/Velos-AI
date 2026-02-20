"""
Velos - Hybrid Resume Parser
Implements a smart parsing pipeline that automatically switches between
text extraction and OCR based on content quality.

Built for ZYND AIckathon 2025
"""

import re
import io
from pathlib import Path
from typing import Optional, Tuple, BinaryIO, Union

# PDF handling
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ pdfplumber not installed. PDF parsing disabled.")

# DOCX handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOCX parsing disabled.")

# OCR handling
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ pytesseract/Pillow not installed. OCR disabled.")

# PDF to Image for OCR
try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("⚠️ pdf2image not installed. PDF OCR fallback disabled.")


class ResumeParser:
    """
    Hybrid Resume Parser with automatic OCR fallback.
    
    Strategy:
    1. Attempt fast text extraction first
    2. If text is insufficient (<50 chars), switch to OCR
    3. Clean and sanitize the extracted text
    """
    
    # Minimum characters to consider text extraction successful
    MIN_TEXT_THRESHOLD = 50
    
    def __init__(self, ocr_lang: str = "eng"):
        """
        Initialize the parser.
        
        Args:
            ocr_lang: Language for OCR (default: English)
        """
        self.ocr_lang = ocr_lang
        self.last_method_used = None  # Track which method was used
        
    def parse_file(self, file_data: Union[bytes, BinaryIO], 
                   filename: str) -> Tuple[str, dict]:
        """
        Main entry point - parse a resume file.
        
        Args:
            file_data: File content as bytes or file-like object
            filename: Original filename (used to determine type)
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        # Convert to bytes if file-like object
        if hasattr(file_data, 'read'):
            file_bytes = file_data.read()
        else:
            file_bytes = file_data
            
        # Get file extension
        ext = Path(filename).suffix.lower()
        
        metadata = {
            "filename": filename,
            "extension": ext,
            "size_bytes": len(file_bytes),
            "extraction_method": None,
            "ocr_used": False,
            "pages_processed": 0,
            "warnings": []
        }
        
        # Route to appropriate handler
        if ext == ".pdf":
            if not PDF_AVAILABLE:
                raise ValueError("PDF parsing not available. Install pdfplumber.")
            text, metadata = self._parse_pdf(file_bytes, metadata)
            
        elif ext == ".docx":
            if not DOCX_AVAILABLE:
                raise ValueError("DOCX parsing not available. Install python-docx.")
            text, metadata = self._parse_docx(file_bytes, metadata)
            
        elif ext == ".doc":
            metadata["warnings"].append("Legacy .doc format - limited support")
            # Try to extract as plain text
            try:
                text = file_bytes.decode('utf-8', errors='ignore')
            except Exception:
                text = file_bytes.decode('latin-1', errors='ignore')
            metadata["extraction_method"] = "raw_decode"
            
        elif ext == ".txt":
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1', errors='ignore')
            metadata["extraction_method"] = "plain_text"
            
        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            if not OCR_AVAILABLE:
                raise ValueError("Image OCR not available. Install pytesseract and Pillow.")
            text, metadata = self._parse_image(file_bytes, metadata)
            
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Sanitize the text
        text = self._sanitize_text(text)
        metadata["char_count"] = len(text)
        metadata["word_count"] = len(text.split())
        
        self.last_method_used = metadata["extraction_method"]
        
        return text, metadata
    
    def _parse_pdf(self, file_bytes: bytes, metadata: dict) -> Tuple[str, dict]:
        """
        Parse PDF with hybrid text/OCR approach.
        
        Strategy:
        1. Try text extraction with pdfplumber
        2. If page has < 50 chars, use OCR for that page
        """
        all_text = []
        ocr_pages = []
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            metadata["pages_processed"] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Try text extraction first
                page_text = page.extract_text() or ""
                
                # Check if text extraction was successful
                if len(page_text.strip()) >= self.MIN_TEXT_THRESHOLD:
                    all_text.append(page_text)
                else:
                    # Text extraction failed - try OCR
                    ocr_pages.append(page_num)
                    
                    if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                        try:
                            # Convert page to image
                            page_image = page.to_image(resolution=300)
                            pil_image = page_image.original
                            
                            # Run OCR
                            ocr_text = pytesseract.image_to_string(
                                pil_image, 
                                lang=self.ocr_lang
                            )
                            all_text.append(ocr_text)
                            metadata["ocr_used"] = True
                        except Exception as e:
                            metadata["warnings"].append(
                                f"OCR failed on page {page_num}: {str(e)}"
                            )
                            # Still append whatever text we got
                            all_text.append(page_text)
                    else:
                        # OCR not available, use whatever we got
                        all_text.append(page_text)
                        metadata["warnings"].append(
                            f"Page {page_num} appears to be scanned but OCR is not available"
                        )
        
        # Determine extraction method
        if metadata["ocr_used"]:
            metadata["extraction_method"] = "hybrid_text_ocr"
            metadata["ocr_pages"] = ocr_pages
        else:
            metadata["extraction_method"] = "text_extraction"
        
        return "\n\n".join(all_text), metadata
    
    def _parse_pdf_ocr_fallback(self, file_bytes: bytes, metadata: dict) -> Tuple[str, dict]:
        """
        Full OCR fallback for PDFs that completely fail text extraction.
        Uses pdf2image to convert all pages to images.
        """
        if not PDF2IMAGE_AVAILABLE or not OCR_AVAILABLE:
            raise ValueError("Full OCR requires pdf2image and pytesseract")
        
        all_text = []
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(file_bytes, dpi=300)
            metadata["pages_processed"] = len(images)
            
            for page_num, image in enumerate(images, 1):
                try:
                    ocr_text = pytesseract.image_to_string(image, lang=self.ocr_lang)
                    all_text.append(ocr_text)
                except Exception as e:
                    metadata["warnings"].append(f"OCR failed on page {page_num}: {str(e)}")
            
            metadata["extraction_method"] = "full_ocr"
            metadata["ocr_used"] = True
            
        except Exception as e:
            raise ValueError(f"PDF OCR conversion failed: {str(e)}")
        
        return "\n\n".join(all_text), metadata
    
    def _parse_docx(self, file_bytes: bytes, metadata: dict) -> Tuple[str, dict]:
        """
        Parse DOCX file using python-docx.
        """
        doc = Document(io.BytesIO(file_bytes))
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        metadata["extraction_method"] = "docx_extraction"
        metadata["pages_processed"] = 1  # DOCX doesn't have clear page boundaries
        
        return "\n".join(paragraphs), metadata
    
    def _parse_image(self, file_bytes: bytes, metadata: dict) -> Tuple[str, dict]:
        """
        Parse image file using OCR.
        """
        image = Image.open(io.BytesIO(file_bytes))
        
        # Convert to RGB if necessary (for PNG with alpha channel)
        if image.mode in ('RGBA', 'LA', 'P'):
            image = image.convert('RGB')
        
        ocr_text = pytesseract.image_to_string(image, lang=self.ocr_lang)
        
        metadata["extraction_method"] = "image_ocr"
        metadata["ocr_used"] = True
        metadata["pages_processed"] = 1
        metadata["image_size"] = f"{image.width}x{image.height}"
        
        return ocr_text, metadata
    
    def _sanitize_text(self, text: str) -> str:
        """
        Clean and sanitize extracted text.
        
        - Remove null bytes
        - Normalize whitespace
        - Remove control characters
        - Fix common OCR artifacts
        """
        if not text:
            return ""
        
        # Remove null bytes
        text = text.replace('\x00', '')
        
        # Remove other control characters (except newlines and tabs)
        text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Replace multiple spaces with single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace multiple newlines with double newline (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Fix common OCR artifacts (safe fixes only)
        text = text.replace('—', '-').replace('–', '-')
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        text = text.replace('…', '...')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_supported_formats(self) -> dict:
        """
        Return dictionary of supported formats and their availability.
        """
        return {
            ".pdf": PDF_AVAILABLE,
            ".docx": DOCX_AVAILABLE,
            ".doc": True,  # Basic support
            ".txt": True,
            ".png": OCR_AVAILABLE,
            ".jpg": OCR_AVAILABLE,
            ".jpeg": OCR_AVAILABLE,
            ".tiff": OCR_AVAILABLE,
            ".bmp": OCR_AVAILABLE,
            "ocr_fallback": PDF2IMAGE_AVAILABLE and OCR_AVAILABLE
        }


# Singleton instance for easy import
resume_parser = ResumeParser()


def parse_resume(file_data: Union[bytes, BinaryIO], filename: str) -> Tuple[str, dict]:
    """
    Convenience function to parse a resume.
    
    Args:
        file_data: File content
        filename: Original filename
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    return resume_parser.parse_file(file_data, filename)


# Quick test
if __name__ == "__main__":
    parser = ResumeParser()
    print("📄 Resume Parser initialized")
    print(f"Supported formats: {parser.get_supported_formats()}")
